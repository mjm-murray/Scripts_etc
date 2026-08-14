"""Convert a (small subset of) Markdown to a Word .docx.

Supports:
- `# Heading 1` / `## Heading 2` / `### Heading 3`
- `- **Speaker:** body text` (speaker bolded, rest normal)
- `- plain bullet`
- Blank lines as paragraph separators
- Any other line as a normal paragraph

If the input file has no markdown structure (no headings, no bullets), it falls
back to splitting the text into ~3-sentence bulleted chunks so the output is at
least readable.

Usage: python md-to-docx.py <input.md> <output.docx> [--title "..."]
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


SPEAKER_BULLET = re.compile(r'^\s*-\s+\*\*(?P<speaker>[^*]+?):\*\*\s*(?P<body>.*)$')
PLAIN_BULLET   = re.compile(r'^\s*-\s+(?P<body>.+)$')
HEADING        = re.compile(r'^(?P<hashes>#{1,6})\s+(?P<text>.+)$')


def split_sentences(text: str, group: int = 3) -> list[str]:
    text = text.strip()
    if not text:
        return []
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
    if not sentences:
        return [text]
    out: list[str] = []
    chunk: list[str] = []
    for s in sentences:
        chunk.append(s)
        if len(chunk) >= group:
            out.append(' '.join(chunk))
            chunk = []
    if chunk:
        out.append(' '.join(chunk))
    return out


def is_markdown_structured(lines: list[str]) -> bool:
    for ln in lines:
        if HEADING.match(ln) or PLAIN_BULLET.match(ln):
            return True
    return False


def add_bullet(doc: Document, speaker: str | None, body: str) -> None:
    p = doc.add_paragraph(style='List Bullet')
    if speaker:
        run = p.add_run(f'{speaker}: ')
        run.bold = True
    p.add_run(body)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument('input', type=Path)
    ap.add_argument('output', type=Path)
    ap.add_argument('--title', default=None)
    args = ap.parse_args()

    text = args.input.read_text(encoding='utf-8', errors='replace')
    lines = text.splitlines()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    if args.title:
        doc.add_heading(args.title, level=0)

    if not is_markdown_structured(lines):
        for chunk in split_sentences(text):
            add_bullet(doc, None, chunk)
        doc.save(str(args.output))
        print(f'Wrote {args.output} (fallback bullet formatting)')
        return 0

    in_blank = True
    for raw in lines:
        ln = raw.rstrip()
        if not ln.strip():
            in_blank = True
            continue

        m = HEADING.match(ln)
        if m:
            level = min(len(m.group('hashes')), 4)
            doc.add_heading(m.group('text').strip(), level=level)
            in_blank = True
            continue

        m = SPEAKER_BULLET.match(ln)
        if m:
            add_bullet(doc, m.group('speaker').strip(), m.group('body').strip())
            in_blank = False
            continue

        m = PLAIN_BULLET.match(ln)
        if m:
            add_bullet(doc, None, m.group('body').strip())
            in_blank = False
            continue

        doc.add_paragraph(ln.strip())
        in_blank = False

    doc.save(str(args.output))
    print(f'Wrote {args.output}')
    return 0


if __name__ == '__main__':
    sys.exit(main())
