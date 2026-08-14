"""Transcribe Video - portable, single-file tool.

Pick a video/audio file, run a 4-stage pipeline against it, and write the
outputs next to the source:

  1. <name>.transcript.md           - faster-whisper transcript
  2. <name>.transcript.formatted.md - speaker-attributed bullets (needs `claude` CLI)
  3. <name>.transcript.docx         - Word doc (python-docx)
  4. <name>.notes.md                - meeting notes (needs `claude` CLI)

Everything except the optional `claude` CLI step is fully self-contained:
faster-whisper, a bundled ffmpeg (via imageio-ffmpeg), and python-docx are all
embedded in the .exe by PyInstaller. The whisper model is downloaded into a
`models/` folder beside the .exe on first run, so the .exe + models can travel
together on a USB drive.

Run directly with `python transcribe_gui.py` during development, or build into
a single-file .exe via `build.ps1`.
"""
from __future__ import annotations

import datetime as dt
import os
import queue
import re
import shutil
import subprocess
import sys
import threading
import tkinter as tk
import traceback
from pathlib import Path
from tkinter import filedialog, messagebox, scrolledtext, ttk

APP_TITLE = "Transcribe Video"
WHISPER_MODELS = ["turbo", "large-v3", "medium", "small", "base", "tiny"]
VIDEO_TYPES = [
    ("Video/audio files", "*.mp4 *.m4a *.mov *.mkv *.webm *.wav *.mp3 *.flac *.ogg *.aac"),
    ("All files", "*.*"),
]

# faster-whisper exposes "large-v3-turbo" under the short name "turbo" via this alias.
WHISPER_NAME_ALIAS = {"turbo": "large-v3-turbo"}


# ---------------------------------------------------------------------------
# Paths (works in dev and when bundled by PyInstaller --onefile)
# ---------------------------------------------------------------------------

def exe_dir() -> Path:
    """Directory the .exe (or .py during dev) lives in. Persistent across runs."""
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def models_dir() -> Path:
    d = exe_dir() / "models"
    d.mkdir(parents=True, exist_ok=True)
    return d


def hidden_subprocess_kwargs() -> dict:
    if os.name != "nt":
        return {}
    si = subprocess.STARTUPINFO()
    si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
    return {"startupinfo": si, "creationflags": subprocess.CREATE_NO_WINDOW}


# ---------------------------------------------------------------------------
# Pipeline stages (each callable returns the output path or None)
# ---------------------------------------------------------------------------

def stage_whisper(audio: Path, transcript_md: Path, model_name: str, log) -> None:
    """Transcribe audio with faster-whisper. Writes plain text (no punctuation needed)."""
    if transcript_md.exists():
        log(f"Skip whisper (transcript exists): {transcript_md.name}")
        return

    # Make sure faster-whisper can find ffmpeg for any container it can't decode itself.
    # PyAV (bundled by faster-whisper) handles most formats natively, but having ffmpeg on
    # PATH is a safety net.
    try:
        import imageio_ffmpeg
        ff = imageio_ffmpeg.get_ffmpeg_exe()
        ff_dir = str(Path(ff).parent)
        if ff_dir not in os.environ.get("PATH", ""):
            os.environ["PATH"] = ff_dir + os.pathsep + os.environ.get("PATH", "")
        log(f"ffmpeg: {ff}")
    except Exception as e:
        log(f"WARNING: imageio-ffmpeg not available ({e}); relying on system ffmpeg.")

    from faster_whisper import WhisperModel

    repo_name = WHISPER_NAME_ALIAS.get(model_name, model_name)
    cache = models_dir()
    log(f"Loading whisper model '{model_name}' (cache={cache}) ...")
    # Pin CPU/int8 for portability — bundling CUDA DLLs would balloon the .exe and break on
    # machines without an NVIDIA GPU. CTranslate2 int8 on CPU is fast enough for meeting audio.
    model = WhisperModel(repo_name, device="cpu", compute_type="int8", download_root=str(cache))
    log("Transcribing... (whisper is the slow part)")

    segments, info = model.transcribe(str(audio), language="en", vad_filter=False)
    log(f"Detected duration: {getattr(info, 'duration', '?')}s")

    chunks: list[str] = []
    for seg in segments:
        text = (seg.text or "").strip()
        if text:
            chunks.append(text)
            log(f"  [{seg.start:7.2f} -> {seg.end:7.2f}]  {text[:120]}")
    body = " ".join(chunks)
    transcript_md.write_text(body, encoding="utf-8")
    log(f"Wrote {transcript_md.name}  ({len(body):,} chars)")


def have_claude() -> bool:
    return shutil.which("claude") is not None


def run_claude(prompt: str, stdin_text: str, log) -> str | None:
    """Pipe stdin_text into `claude -p --output-format text <prompt>`. Returns stdout text or None on failure."""
    try:
        proc = subprocess.run(
            ["claude", "-p", "--output-format", "text", prompt],
            input=stdin_text,
            text=True,
            capture_output=True,
            **hidden_subprocess_kwargs(),
        )
    except OSError as e:
        log(f"claude launch failed: {e}")
        return None
    if proc.returncode != 0:
        log(f"claude exited with code {proc.returncode}: {proc.stderr.strip()[:300]}")
        return None
    out = (proc.stdout or "").strip()
    if not out:
        log("claude returned empty output.")
        return None
    return out


FORMAT_PROMPT = """\
You are reformatting a raw, unpunctuated meeting transcript provided via stdin. The transcript is from a single audio recording with multiple speakers but no diarization.

Your task:
1. Split the transcript into individual speaker turns based on context cues (people addressing each other by name, topic changes, conversational handoffs like "thanks", "go ahead", question-answer pairs).
2. Identify each speaker by NAME if you can infer it confidently from the context (e.g., one participant says "thanks Frank" - the next speaker is likely Frank; or someone introduces themselves). When unsure, label as "Speaker 1", "Speaker 2", etc., assigned consistently across the whole document. DO NOT guess names you are not confident about.
3. Add minimal punctuation and capitalization so the text reads naturally. DO NOT summarize, paraphrase, or omit any words from the original transcript - every word from the input must appear in the output.
4. Output as a Markdown bulleted list. One bullet per speaker turn. Format each bullet as: `- **Speaker Name:** turn text.`

If a single speaker has a very long uninterrupted stretch, you may split it into multiple consecutive bullets at natural sentence/topic boundaries, all attributed to the same speaker.

Output ONLY the bulleted markdown. No preamble, no headings, no commentary, no code fences.
"""


def stage_format(transcript_md: Path, formatted_md: Path, log) -> None:
    if formatted_md.exists():
        log(f"Skip format (exists): {formatted_md.name}")
        return
    if not have_claude():
        log("claude CLI not on PATH - skipping speaker-attributed formatting.")
        return
    log("Formatting transcript with speaker attribution (claude) ...")
    out = run_claude(FORMAT_PROMPT, transcript_md.read_text(encoding="utf-8", errors="replace"), log)
    if out:
        formatted_md.write_text(out, encoding="utf-8")
        log(f"Wrote {formatted_md.name}  ({len(out):,} chars)")


def stage_docx(source_md: Path, docx_path: Path, title: str, log) -> None:
    if docx_path.exists():
        log(f"Skip docx (exists): {docx_path.name}")
        return
    log(f"Building docx from {source_md.name} ...")
    md_to_docx(source_md, docx_path, title)
    log(f"Wrote {docx_path.name}")


def notes_prompt(date_str: str) -> str:
    return f"""\
You are an executive assistant producing meeting notes from a transcript provided via stdin.

Produce a single Markdown document with these sections (omit any section with no content):

# <inferred meeting title>

**Date:** {date_str}

## Attendees
(Best guesses from names mentioned, with affiliation if mentioned in the transcript.)

## Summary
(3-6 sentences on what was discussed.)

## Key Takeaways
(Bulleted, high-signal insights and conclusions from the discussion. Each bullet is a self-contained statement a reader could grasp without reading the full transcript. Distinct from "Key Decisions": takeaways are what was learned or observed; decisions are specific commitments made.)

## Key Decisions
(Bullet list of specific decisions or commitments made during the meeting.)

## Action Items
(Bullet list, format: "- [ ] <action> - <owner if known> - <due date if known>")

## Risks / Concerns
(Bullet list of risks, blockers, or concerns raised that aren't already captured as action items or open questions.)

## Open Questions
(Anything left unresolved or requiring follow-up.)

Output ONLY the markdown. No preamble, no code fences, no commentary. Use real line breaks between sections.
"""


def stage_notes(source_md: Path, notes_md: Path, date_str: str, log) -> None:
    if notes_md.exists():
        log(f"Skip notes (exists): {notes_md.name}")
        return
    if not have_claude():
        log("claude CLI not on PATH - skipping meeting-notes generation.")
        return
    log("Generating meeting notes (claude) ...")
    out = run_claude(notes_prompt(date_str), source_md.read_text(encoding="utf-8", errors="replace"), log)
    if out:
        notes_md.write_text(out, encoding="utf-8")
        log(f"Wrote {notes_md.name}  ({len(out):,} chars)")


# ---------------------------------------------------------------------------
# Markdown -> DOCX (embedded from the old md-to-docx.py helper)
# ---------------------------------------------------------------------------

SPEAKER_BULLET = re.compile(r'^\s*-\s+\*\*(?P<speaker>[^*]+?):\*\*\s*(?P<body>.*)$')
PLAIN_BULLET   = re.compile(r'^\s*-\s+(?P<body>.+)$')
HEADING        = re.compile(r'^(?P<hashes>#{1,6})\s+(?P<text>.+)$')


def _split_sentences(text: str, group: int = 3) -> list[str]:
    text = text.strip()
    if not text:
        return []
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
    if not sentences:
        return [text]
    out: list[str] = []
    chunk: list[str] = []
    for s in sentences:
        chunk.append(s)
        if len(chunk) >= group:
            out.append(' '.join(chunk))
            chunk = []
    if chunk:
        out.append(' '.join(chunk))
    return out


def _is_markdown_structured(lines: list[str]) -> bool:
    for ln in lines:
        if HEADING.match(ln) or PLAIN_BULLET.match(ln):
            return True
    return False


def md_to_docx(input_md: Path, output_docx: Path, title: str | None = None) -> None:
    from docx import Document
    from docx.shared import Pt

    text = input_md.read_text(encoding='utf-8', errors='replace')
    lines = text.splitlines()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    if title:
        doc.add_heading(title, level=0)

    def add_bullet(speaker: str | None, body: str) -> None:
        p = doc.add_paragraph(style='List Bullet')
        if speaker:
            run = p.add_run(f'{speaker}: ')
            run.bold = True
        p.add_run(body)

    if not _is_markdown_structured(lines):
        for chunk in _split_sentences(text):
            add_bullet(None, chunk)
        doc.save(str(output_docx))
        return

    for raw in lines:
        ln = raw.rstrip()
        if not ln.strip():
            continue
        m = HEADING.match(ln)
        if m:
            level = min(len(m.group('hashes')), 4)
            doc.add_heading(m.group('text').strip(), level=level)
            continue
        m = SPEAKER_BULLET.match(ln)
        if m:
            add_bullet(m.group('speaker').strip(), m.group('body').strip())
            continue
        m = PLAIN_BULLET.match(ln)
        if m:
            add_bullet(None, m.group('body').strip())
            continue
        doc.add_paragraph(ln.strip())

    doc.save(str(output_docx))


# ---------------------------------------------------------------------------
# Tk GUI
# ---------------------------------------------------------------------------

class TranscribeApp:
    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        root.title(APP_TITLE)
        root.geometry("820x600")
        root.minsize(680, 460)

        self.video_path = tk.StringVar()
        self.model = tk.StringVar(value="turbo")
        self.do_format = tk.BooleanVar(value=True)
        self.do_docx = tk.BooleanVar(value=True)
        self.do_notes = tk.BooleanVar(value=True)
        self.status = tk.StringVar(value="Pick a video file to begin.")

        self._log_queue: queue.Queue[str | None] = queue.Queue()
        self._worker: threading.Thread | None = None
        self._cancel_event = threading.Event()

        self._build_ui()
        self._initial_environment_report()
        root.after(80, self._drain_log_queue)

    def _build_ui(self) -> None:
        pad = {"padx": 10, "pady": 6}
        root = self.root

        top = ttk.Frame(root)
        top.pack(fill=tk.X, **pad)
        ttk.Label(top, text="Video file:").grid(row=0, column=0, sticky="w")
        ttk.Entry(top, textvariable=self.video_path).grid(row=0, column=1, sticky="ew", padx=(6, 6))
        ttk.Button(top, text="Browse...", command=self._on_browse).grid(row=0, column=2)
        top.columnconfigure(1, weight=1)

        opts = ttk.LabelFrame(root, text="Options")
        opts.pack(fill=tk.X, **pad)
        ttk.Label(opts, text="Whisper model:").grid(row=0, column=0, sticky="w", padx=8, pady=6)
        ttk.Combobox(opts, textvariable=self.model, values=WHISPER_MODELS, state="readonly", width=14).grid(row=0, column=1, sticky="w", pady=6)
        ttk.Checkbutton(opts, text="Speaker-attributed formatting (requires `claude` CLI)", variable=self.do_format).grid(row=1, column=0, columnspan=3, sticky="w", padx=8)
        ttk.Checkbutton(opts, text="Build Word .docx", variable=self.do_docx).grid(row=2, column=0, columnspan=3, sticky="w", padx=8)
        ttk.Checkbutton(opts, text="Generate meeting notes (requires `claude` CLI)", variable=self.do_notes).grid(row=3, column=0, columnspan=3, sticky="w", padx=8, pady=(0, 6))

        actions = ttk.Frame(root)
        actions.pack(fill=tk.X, **pad)
        self.run_btn = ttk.Button(actions, text="Process Video", command=self._on_run)
        self.run_btn.pack(side=tk.LEFT)
        self.cancel_btn = ttk.Button(actions, text="Cancel", command=self._on_cancel, state=tk.DISABLED)
        self.cancel_btn.pack(side=tk.LEFT, padx=(8, 0))
        self.open_folder_btn = ttk.Button(actions, text="Open output folder", command=self._on_open_folder, state=tk.DISABLED)
        self.open_folder_btn.pack(side=tk.LEFT, padx=(8, 0))

        log_frame = ttk.LabelFrame(root, text="Log")
        log_frame.pack(fill=tk.BOTH, expand=True, **pad)
        self.log = scrolledtext.ScrolledText(log_frame, wrap=tk.WORD, height=18, state=tk.DISABLED, font=("Consolas", 9))
        self.log.pack(fill=tk.BOTH, expand=True, padx=6, pady=6)

        ttk.Label(root, textvariable=self.status, anchor="w", relief="sunken").pack(fill=tk.X, side=tk.BOTTOM)

    def _initial_environment_report(self) -> None:
        self._log(f"App dir: {exe_dir()}")
        self._log(f"Model cache: {models_dir()}")
        self._log(f"`claude` CLI: {'found at ' + shutil.which('claude') if have_claude() else 'NOT FOUND on PATH (formatting + notes stages will be skipped)'}")
        try:
            import imageio_ffmpeg
            self._log(f"Bundled ffmpeg: {imageio_ffmpeg.get_ffmpeg_exe()}")
        except Exception as e:
            self._log(f"WARNING: bundled ffmpeg not available: {e}")

    def _on_browse(self) -> None:
        initial = ""
        if self.video_path.get():
            initial = str(Path(self.video_path.get()).parent)
        path = filedialog.askopenfilename(title="Pick a video file", initialdir=initial or None, filetypes=VIDEO_TYPES)
        if path:
            self.video_path.set(path)
            self.status.set(f"Selected: {Path(path).name}")

    def _on_open_folder(self) -> None:
        p = self.video_path.get()
        if not p:
            return
        folder = str(Path(p).parent)
        try:
            os.startfile(folder)  # type: ignore[attr-defined]
        except OSError as e:
            messagebox.showerror(APP_TITLE, f"Could not open folder:\n{e}")

    def _on_cancel(self) -> None:
        self._cancel_event.set()
        self.status.set("Cancellation requested — will stop after the current stage.")

    def _on_run(self) -> None:
        path = self.video_path.get().strip()
        if not path:
            messagebox.showwarning(APP_TITLE, "Pick a video file first.")
            return
        video = Path(path)
        if not video.is_file():
            messagebox.showerror(APP_TITLE, f"File not found:\n{path}")
            return

        self._cancel_event.clear()
        self._set_running(True)
        self._log("=" * 78)
        self._log(f"Processing: {video.name}")
        self._log(f"Output folder: {video.parent}")
        self._log(f"Model: {self.model.get()}  Format={self.do_format.get()}  Docx={self.do_docx.get()}  Notes={self.do_notes.get()}")
        self._log("=" * 78)

        self._worker = threading.Thread(target=self._run_pipeline, args=(video,), daemon=True)
        self._worker.start()

    def _run_pipeline(self, video: Path) -> None:
        def log(msg: str) -> None:
            self._log_queue.put(msg)

        try:
            stem = video.stem
            d = video.parent
            transcript = d / f"{stem}.transcript.md"
            formatted  = d / f"{stem}.transcript.formatted.md"
            docx_path  = d / f"{stem}.transcript.docx"
            notes      = d / f"{stem}.notes.md"

            # Stage 1: whisper
            stage_whisper(video, transcript, self.model.get(), log)
            if self._cancel_event.is_set():
                log("Cancelled.")
                return

            # Stage 2: speaker formatting (optional)
            if self.do_format.get():
                stage_format(transcript, formatted, log)
                if self._cancel_event.is_set():
                    log("Cancelled.")
                    return

            # Stage 3: docx (prefer formatted source if present)
            if self.do_docx.get():
                src = formatted if formatted.exists() else transcript
                title = re.sub(r'[_-]+', ' ', stem)
                stage_docx(src, docx_path, title, log)
                if self._cancel_event.is_set():
                    log("Cancelled.")
                    return

            # Stage 4: meeting notes (optional)
            if self.do_notes.get():
                src = formatted if formatted.exists() else transcript
                date_str = dt.datetime.fromtimestamp(video.stat().st_mtime).strftime('%Y-%m-%d')
                stage_notes(src, notes, date_str, log)

            log("")
            log("All stages complete.")
        except Exception:
            log("PIPELINE ERROR:")
            for line in traceback.format_exc().splitlines():
                log(f"  {line}")
        finally:
            self._log_queue.put(None)  # sentinel

    def _drain_log_queue(self) -> None:
        try:
            while True:
                item = self._log_queue.get_nowait()
                if item is None:
                    self._set_running(False)
                    self._summarize_outputs()
                else:
                    self._log(item)
        except queue.Empty:
            pass
        self.root.after(80, self._drain_log_queue)

    def _summarize_outputs(self) -> None:
        video = Path(self.video_path.get())
        if not video.exists():
            self.status.set("Done.")
            return
        stem = video.stem
        d = video.parent
        artifacts = {
            "transcript": d / f"{stem}.transcript.md",
            "formatted":  d / f"{stem}.transcript.formatted.md",
            "docx":       d / f"{stem}.transcript.docx",
            "notes":      d / f"{stem}.notes.md",
        }
        present = {k: v for k, v in artifacts.items() if v.exists()}
        self._log("")
        self._log("Output files:")
        for label, p in present.items():
            kb = p.stat().st_size / 1024
            self._log(f"  [{label:10}] {p.name}  ({kb:,.1f} KB)")
        missing = [k for k in artifacts if k not in present]
        if missing:
            self._log(f"  (not produced: {', '.join(missing)})")
        self.status.set(f"Done. {len(present)}/{len(artifacts)} artifacts in {d}")
        self.open_folder_btn.configure(state=tk.NORMAL)

    def _set_running(self, running: bool) -> None:
        self.run_btn.configure(state=tk.DISABLED if running else tk.NORMAL)
        self.cancel_btn.configure(state=tk.NORMAL if running else tk.DISABLED)
        if running:
            self.status.set("Running... (whisper is the slow part — first run downloads the model)")

    def _log(self, msg: str) -> None:
        self.log.configure(state=tk.NORMAL)
        self.log.insert(tk.END, msg + "\n")
        self.log.see(tk.END)
        self.log.configure(state=tk.DISABLED)


def main() -> None:
    root = tk.Tk()
    try:
        style = ttk.Style()
        if "vista" in style.theme_names():
            style.theme_use("vista")
    except tk.TclError:
        pass
    TranscribeApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
